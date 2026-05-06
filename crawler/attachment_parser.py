from __future__ import annotations

from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt"}


def parse_attachment(path: str, config: dict[str, Any] | None = None) -> str | None:
    config = config or {}
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return None

    max_chars = int(config.get("max_chars", 200000))
    try:
        if suffix == ".pdf":
            text = _parse_pdf(file_path)
        elif suffix == ".docx":
            text = _parse_docx(file_path)
        elif suffix in {".xlsx", ".xlsm"}:
            text = _parse_excel(file_path)
        elif suffix == ".csv":
            text = file_path.read_text(encoding=config.get("encoding", "utf-8"), errors="ignore")
        else:
            text = file_path.read_text(encoding=config.get("encoding", "utf-8"), errors="ignore")
    except ImportError as exc:
        raise RuntimeError(f"Missing dependency for parsing {suffix}: {exc}") from exc

    text = text.strip()
    if not text:
        return None
    return text[:max_chars]


def write_parsed_text(path: str, text: str) -> str:
    output_path = Path(path).with_suffix(Path(path).suffix + ".txt")
    output_path.write_text(text, encoding="utf-8")
    return str(output_path)


def _parse_pdf(path: Path) -> str:
    try:
        import pdfplumber

        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n\n".join(page for page in pages if page)
    except ImportError:
        pass

    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError as exc:
        raise ImportError("install pdfplumber or PyPDF2 to parse PDF files") from exc


def _parse_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines = []
    for table in document.tables:
        for row in table.rows:
            table_lines.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs + table_lines)


def _parse_excel(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                lines.append("\t".join(values))
    workbook.close()
    return "\n".join(lines)
